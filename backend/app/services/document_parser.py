import os
from typing import Tuple, Optional, List
import base64
from pathlib import Path
from docx import Document
import fitz  # PyMuPDF
from PIL import Image
import pytesseract


class DocumentParser:
    """Parses various document types and extracts text, code, and images."""

    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".py", ".java", ".cpp", ".js", ".ts", ".png", ".jpg", ".jpeg"}

    @staticmethod
    def parse_document(file_path: str) -> Tuple[str, Optional[str], Optional[List[str]], str]:
        """
        Parse document and extract text, code, and images.
        
        Returns:
            Tuple of (text, code, images_base64, file_type)
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == ".pdf":
            return DocumentParser._parse_pdf(file_path)
        elif file_ext == ".docx":
            return DocumentParser._parse_docx(file_path)
        elif file_ext in {".png", ".jpg", ".jpeg"}:
            return DocumentParser._parse_image(file_path)
        elif file_ext in {".py", ".java", ".cpp", ".js", ".ts", ".txt"}:
            return DocumentParser._parse_code(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

    @staticmethod
    def _parse_pdf(file_path: str) -> Tuple[str, Optional[str], Optional[List[str]], str]:
        """Extract text and images from PDF."""
        text = ""
        images_base64 = []
        
        try:
            doc = fitz.open(file_path)
            
            for page_num, page in enumerate(doc):
                # Extract text
                text += f"--- Page {page_num + 1} ---\n"
                text += page.get_text()
                text += "\n"
                
                # Extract images
                image_list = page.get_images()
                for img_index in image_list:
                    try:
                        xref = img_index[0]  # First element is the xref integer
                        pix = fitz.Pixmap(doc, xref)
                        
                        # Convert CMYK or other color spaces to RGB
                        if pix.n >= 5:
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                            
                        img_data = pix.tobytes("png")
                        img_base64 = base64.b64encode(img_data).decode("utf-8")
                        images_base64.append(f"data:image/png;base64,{img_base64}")
                        pix = None
                    except Exception:
                        continue
            
            doc.close()
            return text, None, images_base64 if images_base64 else None, "pdf"
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")

    @staticmethod
    def _parse_docx(file_path: str) -> Tuple[str, Optional[str], Optional[List[str]], str]:
        """Extract text and images from DOCX."""
        text = ""
        images_base64 = []
        
        try:
            doc = Document(file_path)
            
            # Extract text
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    text += row_text + "\n"
            
            # Extract images from document relationships
            try:
                for rel in doc.part.rels.values():
                    if hasattr(rel, "target_ref") and "image" in str(rel.target_ref):
                        if not getattr(rel, "is_external", False):
                            image_data = rel.target_part.blob
                            img_base64 = base64.b64encode(image_data).decode("utf-8")
                            images_base64.append(f"data:image/png;base64,{img_base64}")
            except Exception:
                pass
            
            return text, None, images_base64 if images_base64 else None, "docx"
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")

    @staticmethod
    def _parse_image(file_path: str) -> Tuple[str, Optional[str], Optional[List[str]], str]:
        """Extract text from image using OCR."""
        text = ""
        images_base64 = []
        
        try:
            img = Image.open(file_path)
            
            # Use Tesseract OCR with fallback
            try:
                text = pytesseract.image_to_string(img)
            except Exception:
                text = "[OCR text extraction unavailable: Tesseract engine not found on server]"
            
            # Encode image as base64
            with open(file_path, "rb") as f:
                img_data = f.read()
                img_base64 = base64.b64encode(img_data).decode("utf-8")
                images_base64 = [f"data:image/png;base64,{img_base64}"]
            
            return text, None, images_base64, "image"
        except Exception as e:
            raise ValueError(f"Error parsing image: {str(e)}")

    @staticmethod
    def _parse_code(file_path: str) -> Tuple[str, Optional[str], Optional[List[str]], str]:
        """Extract code from code files."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                code = f.read()
            return code, code, None, "code"
        except Exception as e:
            raise ValueError(f"Error parsing code file: {str(e)}")
